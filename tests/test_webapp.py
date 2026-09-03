import os
import time

import pytest
from fastapi.testclient import TestClient


@pytest.fixture()
def client(tmp_path, monkeypatch):
    monkeypatch.setenv("INSIGHT_ENGINE_LLM", "mock")
    monkeypatch.setenv("INSIGHT_ENGINE_DB", str(tmp_path / "insights.db"))

    import insight_engine.webapp as webapp

    webapp._store = None
    return TestClient(webapp.app)


def test_interviews_page_empty(client):
    resp = client.get("/interviews")
    assert resp.status_code == 200
    assert "Загрузите запись" in resp.text


def test_no_auth_required_when_users_not_configured(client):
    resp = client.get("/interviews")
    assert resp.status_code == 200


def test_auth_required_once_users_are_configured(tmp_path, monkeypatch):
    monkeypatch.setenv("INSIGHT_ENGINE_LLM", "mock")
    monkeypatch.setenv("INSIGHT_ENGINE_DB", str(tmp_path / "insights.db"))
    monkeypatch.setenv("INSIGHT_ENGINE_USERS", "alisher:s3cret,farida:hunter2")

    import insight_engine.webapp as webapp

    webapp._store = None
    client = TestClient(webapp.app)

    resp = client.get("/interviews")
    assert resp.status_code == 401

    resp = client.get("/interviews", auth=("alisher", "wrong-password"))
    assert resp.status_code == 401

    resp = client.get("/interviews", auth=("alisher", "s3cret"))
    assert resp.status_code == 200

    resp = client.get("/interviews", auth=("farida", "hunter2"))
    assert resp.status_code == 200


def test_upload_text_transcript_end_to_end(client):
    transcript = (
        "[00:05] Интервьюер: Расскажите про выписки.\n"
        "[00:12] Фарида Р.: Приложение отдаёт выписку только за месяц целиком, "
        "а для отчётности мне нужен точный диапазон дат, это очень неудобно.\n"
    )
    resp = client.post(
        "/upload",
        data={
            "title": "Интервью тест",
            "respondent": "Фарида Р.",
            "segment_label": "Малый бизнес",
            "source": "transcript_text",
            "transcript_text": transcript,
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    job_url = resp.headers["location"]
    assert job_url.startswith("/processing/")
    job_id = job_url.rsplit("/", 1)[-1]

    for _ in range(50):
        status_resp = client.get(f"/api/jobs/{job_id}")
        payload = status_resp.json()
        if payload["status"] in ("done", "error"):
            break
        time.sleep(0.05)
    else:
        pytest.fail("job did not finish in time")

    assert payload["status"] == "done", payload
    interview_id = payload["interview_id"]

    detail = client.get(f"/interviews/{interview_id}")
    assert detail.status_code == 200
    assert "Фарида Р." in detail.text
    assert "insight-card" in detail.text

    hyps = client.get("/hypotheses")
    assert hyps.status_code == 200
    assert "Гипотезы" in hyps.text


def test_upload_docx_transcript_end_to_end(client):
    import io

    from docx import Document

    document = Document()
    document.add_paragraph("[00:05] Интервьюер: Расскажите про выписки.")
    document.add_paragraph(
        "[00:12] Фарида Р.: Приложение отдаёт выписку только за месяц целиком, "
        "а для отчётности мне нужен точный диапазон дат, это очень неудобно."
    )
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    resp = client.post(
        "/upload",
        data={
            "title": "Интервью docx",
            "respondent": "Фарида Р.",
            "segment_label": "Малый бизнес",
            "source": "transcript_file",
        },
        files={
            "transcript_file": (
                "interview.docx",
                buffer,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    job_id = resp.headers["location"].rsplit("/", 1)[-1]

    for _ in range(50):
        status_resp = client.get(f"/api/jobs/{job_id}")
        payload = status_resp.json()
        if payload["status"] in ("done", "error"):
            break
        time.sleep(0.05)
    else:
        pytest.fail("job did not finish in time")

    assert payload["status"] == "done", payload


def test_upload_rejects_empty_text(client):
    resp = client.post(
        "/upload",
        data={
            "title": "X", "respondent": "Y", "segment_label": "Розничный клиент",
            "source": "transcript_text", "transcript_text": "   ",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert "msg=" in resp.headers["location"]
